from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

# Crear un nuevo documento
doc = Document()

# Configurar estilos
styles = doc.styles

# Estilo para título principal
style_title = styles.add_style('TituloPrincipal', WD_STYLE_TYPE.PARAGRAPH)
style_title.font.name = 'Calibri'
style_title.font.size = Pt(24)
style_title.font.bold = True
style_title.font.color.rgb = RGBColor(0, 112, 192)

# Estilo para subtítulos
style_heading1 = styles.add_style('Subtitulo1', WD_STYLE_TYPE.PARAGRAPH)
style_heading1.font.name = 'Calibri'
style_heading1.font.size = Pt(18)
style_heading1.font.bold = True
style_heading1.font.color.rgb = RGBColor(0, 112, 192)

# Estilo para subtítulos secundarios
style_heading2 = styles.add_style('Subtitulo2', WD_STYLE_TYPE.PARAGRAPH)
style_heading2.font.name = 'Calibri'
style_heading2.font.size = Pt(14)
style_heading2.font.bold = True
style_heading2.font.color.rgb = RGBColor(79, 129, 189)

# Estilo para texto normal
style_normal = styles.add_style('TextoNormal', WD_STYLE_TYPE.PARAGRAPH)
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

# Estilo para código
style_code = styles.add_style('Codigo', WD_STYLE_TYPE.PARAGRAPH)
style_code.font.name = 'Consolas'
style_code.font.size = Pt(10)
style_code.font.color.rgb = RGBColor(0, 0, 0)

# Título principal
title = doc.add_paragraph('Documentación: Azure Storage App', style='TituloPrincipal')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Agregar descripción general
doc.add_paragraph('Aplicación en Python con APIs asíncronas para gestionar la descarga y carga de archivos a Azure Storage.', style='TextoNormal')

# Tabla de contenidos
doc.add_paragraph('Tabla de Contenidos', style='Subtitulo1')
toc_items = [
    "1. Análisis de Requerimientos",
    "2. Arquitectura de la Aplicación",
    "3. Guía de Instalación y Configuración",
    "4. Documentación de la API",
    "5. Guía de Usuario",
    "6. Estructura del Proyecto"
]
for item in toc_items:
    p = doc.add_paragraph(style='TextoNormal')
    p.add_run(item)

# Separador
doc.add_paragraph()

# 1. Análisis de Requerimientos
doc.add_paragraph('1. Análisis de Requerimientos', style='Subtitulo1')

doc.add_paragraph('1.1 Requerimientos Funcionales', style='Subtitulo2')
req_func = [
    "Gestión de Archivos en Azure Storage:",
    "- Carga (upload) de archivos a Azure Blob Storage",
    "- Descarga (download) de archivos desde Azure Blob Storage",
    "- Listado de archivos disponibles en un contenedor",
    "- Eliminación de archivos de Azure Blob Storage",
    "",
    "API Asíncrona:",
    "- Implementación de endpoints asíncronos usando FastAPI",
    "- Manejo eficiente de operaciones de larga duración",
    "- Soporte para múltiples operaciones concurrentes",
    "",
    "Autenticación y Seguridad:",
    "- Conexión segura a Azure Storage mediante credenciales",
    "- Validación de permisos para operaciones de archivos",
    "- Manejo seguro de información sensible (claves, tokens)",
    "",
    "Manejo de Errores:",
    "- Gestión adecuada de excepciones durante operaciones de almacenamiento",
    "- Respuestas de error informativas y útiles",
    "- Registro (logging) de errores para diagnóstico"
]
for item in req_func:
    p = doc.add_paragraph(style='TextoNormal')
    p.add_run(item)

doc.add_paragraph('1.2 Especificaciones Técnicas', style='Subtitulo2')
req_tech = [
    "Tecnologías:",
    "- Python 3.x",
    "- FastAPI para APIs asíncronas",
    "- Azure Storage Blob SDK para Python",
    "- Uvicorn como servidor ASGI",
    "",
    "Estructura de Componentes:",
    "- Módulo de configuración para Azure Storage",
    "- Módulo de servicios para operaciones de almacenamiento",
    "- Módulo de API para endpoints HTTP",
    "- Módulo de utilidades para funciones auxiliares",
    "- Módulo de modelos para definición de datos",
    "",
    "Endpoints de API:",
    "- POST /api/files/upload: Carga de archivos",
    "- GET /api/files/download/{filename}: Descarga de archivos",
    "- GET /api/files/list: Listado de archivos",
    "- DELETE /api/files/{filename}: Eliminación de archivos",
    "",
    "Parámetros y Respuestas:",
    "- Carga: Recibe archivo multipart, devuelve URL y metadatos",
    "- Descarga: Recibe nombre de archivo, devuelve stream de datos",
    "- Listado: Recibe parámetros opcionales de filtrado, devuelve lista de archivos",
    "- Eliminación: Recibe nombre de archivo, devuelve confirmación",
    "",
    "Configuración:",
    "- Variables de entorno para credenciales de Azure",
    "- Archivo de configuración para parámetros de la aplicación",
    "- Soporte para diferentes entornos (desarrollo, producción)"
]
for item in req_tech:
    p = doc.add_paragraph(style='TextoNormal')
    p.add_run(item)

# 2. Arquitectura de la Aplicación
doc.add_paragraph('2. Arquitectura de la Aplicación', style='Subtitulo1')

doc.add_paragraph('2.1 Estructura del Proyecto', style='Subtitulo2')
p = doc.add_paragraph(style='Codigo')
p.add_run("""
azure_storage_app/
│
├── src/                          # Código fuente de la aplicación
│   ├── __init__.py               # Inicializador del paquete
│   ├── main.py                   # Punto de entrada de la aplicación
│   ├── config/                   # Configuración de la aplicación
│   │   ├── __init__.py
│   │   ├── settings.py           # Configuración general y variables de entorno
│   │   └── azure_config.py       # Configuración específica de Azure Storage
│   │
│   ├── models/                   # Modelos de datos y esquemas
│   │   ├── __init__.py
│   │   └── schemas.py            # Esquemas Pydantic para validación de datos
│   │
│   ├── services/                 # Servicios de negocio
│   │   ├── __init__.py
│   │   └── storage_service.py    # Servicio para operaciones con Azure Storage
│   │
│   ├── api/                      # Endpoints de la API
│   │   ├── __init__.py
│   │   ├── routes.py             # Registro de rutas
│   │   └── endpoints/            # Endpoints específicos
│   │       ├── __init__.py
│   │       └── files.py          # Endpoints para gestión de archivos
│   │
│   └── utils/                    # Utilidades y helpers
│       ├── __init__.py
│       ├── exceptions.py         # Excepciones personalizadas
│       └── helpers.py            # Funciones auxiliares
│
├── tests/                        # Pruebas unitarias y de integración
│   ├── __init__.py
│   ├── test_api.py               # Pruebas para la API
│   └── test_services.py          # Pruebas para los servicios
│
├── docs/                         # Documentación
│   ├── requerimientos.md         # Análisis de requerimientos
│   ├── arquitectura.md           # Diseño de arquitectura
│   ├── api.md                    # Documentación de la API
│   └── guia_usuario.md           # Guía de usuario
│
├── .env.example                  # Ejemplo de variables de entorno
├── requirements.txt              # Dependencias del proyecto
└── README.md                     # Documentación principal
""")

doc.add_paragraph('2.2 Descripción de Componentes', style='Subtitulo2')
components = [
    "Configuración (src/config/):",
    "- settings.py: Gestiona la configuración general de la aplicación, carga variables de entorno y define parámetros globales.",
    "- azure_config.py: Contiene la configuración específica para Azure Storage, incluyendo la inicialización de clientes y la gestión de credenciales.",
    "",
    "Modelos (src/models/):",
    "- schemas.py: Define los esquemas Pydantic para validación de datos de entrada y salida, asegurando la integridad de los datos en la API.",
    "",
    "Servicios (src/services/):",
    "- storage_service.py: Implementa la lógica de negocio para interactuar con Azure Storage, encapsulando las operaciones de carga, descarga, listado y eliminación de archivos.",
    "",
    "API (src/api/):",
    "- routes.py: Registra y configura las rutas de la API.",
    "- endpoints/files.py: Implementa los endpoints específicos para la gestión de archivos, utilizando los servicios correspondientes.",
    "",
    "Utilidades (src/utils/):",
    "- exceptions.py: Define excepciones personalizadas para manejar errores específicos de la aplicación.",
    "- helpers.py: Proporciona funciones auxiliares reutilizables en diferentes partes de la aplicación."
]
for item in components:
    p = doc.add_paragraph(style='TextoNormal')
    p.add_run(item)

doc.add_paragraph('2.3 Flujo de Datos', style='Subtitulo2')
flow = [
    "1. Las solicitudes HTTP llegan a los endpoints definidos en api/endpoints/.",
    "2. Los endpoints validan los datos de entrada utilizando los esquemas de models/.",
    "3. Los endpoints llaman a los métodos correspondientes en services/.",
    "4. Los servicios utilizan la configuración de config/ para conectarse a Azure Storage.",
    "5. Los servicios realizan las operaciones solicitadas y devuelven los resultados.",
    "6. Los endpoints transforman los resultados en respuestas HTTP adecuadas."
]
for item in flow:
    p = doc.add_paragraph(style='TextoNormal')
    p.add_run(item)

doc.add_paragraph('2.4 Principios de Diseño', style='Subtitulo2')
principles = [
    "1. Separación de Responsabilidades: Cada componente tiene una responsabilidad única y bien definida.",
    "2. Inyección de Dependencias: Los servicios reciben sus dependencias (como configuración) en lugar de crearlas internamente.",
    "3. Asincronía: Todas las operaciones de E/S se implementan de forma asíncrona para maximizar el rendimiento.",
    "4. Manejo de Errores Centralizado: Las excepciones se capturan y procesan de manera consistente en toda la aplicación.",
    "5. Configuración Externalizada: Los parámetros de configuración se almacenan fuera del código para facilitar su gestión."
]
for item in principles:
    p = doc.add_paragraph(style='TextoNormal')
    p.add_run(item)

# 3. Guía de Instalación y Configuración
doc.add_paragraph('3. Guía de Instalación y Configuración', style='Subtitulo1')

doc.add_paragraph('3.1 Requisitos', style='Subtitulo2')
requirements = [
    "- Python 3.7 o superior",
    "- pip (gestor de paquetes de Python)",
    "- Cuenta de Azure con servicio de Azure Storage"
]
for item in requirements:
    p = doc.add_paragraph(style='TextoNormal')
    p.add_run(item)

doc.add_paragraph('3.2 Instalación', style='Subtitulo2')
installation = [
    "1. Clone o descargue el repositorio:",
    "",
    "   git clone <url-del-repositorio>",
    "   cd azure_storage_app",
    "",
    "2. Cree un entorno virtual e instale las dependencias:",
    "",
    "   python -m venv venv",
    "   source venv/bin/activate  # En Windows: venv\\Scripts\\activate",
    "   pip install -r requirements.txt"
]
for item in installation:
    if item.startswith("   "):
        p = doc.add_paragraph(style='Codigo')
        p.add_run(item.strip())
    else:
        p = doc.add_paragraph(style='TextoNormal')
        p.add_run(item)

doc.add_paragraph('3.3 Configuración', style='Subtitulo2')
configuration = [
    "Para utilizar la aplicación, necesita configurar sus credenciales de Azure Storage. Tiene dos opciones:",
    "",
    "Opción 1: Connection String",
    "1. En el portal de Azure, vaya a su cuenta de Storage",
    "2. En la sección \"Configuración\", seleccione \"Claves de acceso\"",
    "3. Copie el valor de \"Connection string\"",
    "",
    "Opción 2: Nombre de cuenta y clave",
    "1. En el portal de Azure, vaya a su cuenta de Storage",
    "2. En la sección \"Configuración\", seleccione \"Claves de acceso\"",
    "3. Copie el nombre de la cuenta y una de las claves",
    "",
    "Cree un archivo .env en la raíz del proyecto basado en el archivo .env.example:",
    "",
    "   cp .env.example .env",
    "",
    "Edite el archivo .env con sus credenciales:",
    "",
    "   # Opción 1: Connection String",
    "   AZURE_STORAGE_CONNECTION_STRING=su_connection_string",
    "",
    "   # O Opción 2: Nombre de cuenta y clave",
    "   AZURE_STORAGE_ACCOUNT_NAME=su_account_name",
    "   AZURE_STORAGE_ACCOUNT_KEY=su_account_key",
    "",
    "   # Nombre del contenedor (por defecto: \"files\")",
    "   AZURE_STORAGE_CONTAINER_NAME=su_container_name"
]
for item in configuration:
    if item.startswith("   "):
        p = doc.add_paragraph(style='Codigo')
        p.add_run(item.strip())
    else:
        p = doc.add_paragraph(style='TextoNormal')
        p.add_run(item)

doc.add_paragraph('3.4 Ejecución', style='Subtitulo2')
execution = [
    "Para iniciar la aplicación en modo desarrollo:",
    "",
    "   uvicorn src.main:app --reload --host 0.0.0.0 --port 8000",
    "",
    "Para iniciar la aplicación en modo producción:",
    "",
    "   uvicorn src.main:app --host 0.0.0.0 --port 8000",
    "",
    "La API estará disponible en http://localhost:8000 y la documentación en http://localhost:8000/docs."
]
for item in execution:
    if item.startswith("   "):
        p = doc.add_paragraph(style='Codigo')
        p.add_run(item.strip())
    else:
        p = doc.add_paragraph(style='TextoNormal')
        p.add_run(item)

# 4. Documentación de la API
doc.add_paragraph('4. Documentación de la API', style='Subtitulo1')

doc.add_paragraph('4.1 Base URL', style='Subtitulo2')
p = doc.add_paragraph(style='Codigo')
p.add_run("http://localhost:8000/api")

doc.add_paragraph('4.2 Endpoints', style='Subtitulo2')

# Endpoint: Carga de Archivos
doc.add_paragraph('4.2.1 Carga de Archivos', style='Subtitulo2')
upload_endpoint = [
    "Endpoint: POST /files/upload",
    "",
    "Descripción: Carga un archivo a Azure Storage.",
    "",
    "Parámetros:",
    "- file (form-data): Archivo a cargar",
    "",
    "Respuesta exitosa (201 Created):",
    "",
    "   {",
    "     \"file_name\": \"a1b2c3d4_ejemplo.pdf\",",
    "     \"file_size\": 12345,",
    "     \"content_type\": \"application/pdf\",",
    "     \"url\": \"https://cuenta.blob.core.windows.net/contenedor/a1b2c3d4_ejemplo.pdf\"",
    "   }",
    "",
    "Errores posibles:",
    "- 400 Bad Request: No se ha proporcionado ningún archivo",
    "- 500 Internal Server Error: Error al cargar el archivo"
]
for item in upload_endpoint:
    if item.startswith("   "):
        p = doc.add_paragraph(style='Codigo')
        p.add_run(item.strip())
    else:
        p = doc.add_paragraph(style='TextoNormal')
        p.add_run(item)

# Endpoint: Descarga de Archivos
doc.add_paragraph('4.2.2 Descarga de Archivos', style='Subtitulo2')
download_endpoint = [
    "Endpoint: GET /files/download/{file_name}",
    "",
    "Descripción: Descarga un archivo desde Azure Storage.",
    "",
    "Parámetros:",
    "- file_name (path): Nombre del archivo a descargar",
    "",
    "Respuesta exitosa:",
    "- Stream del contenido del archivo con los headers adecuados para descarga",
    "",
    "Errores posibles:",
    "- 404 Not Found: El archivo no existe",
    "- 500 Internal Server Error: Error al descargar el archivo"
]
for item in download_endpoint:
    p = doc.add_paragraph(style='TextoNormal')
    p.add_run(item)

# Endpoint: Listado de Archivos
doc.add_paragraph('4.2.3 Listado de Archivos', style='Subtitulo2')
list_endpoint = [
    "Endpoint: GET /files/list",
    "",
    "Descripción: Lista los archivos disponibles en Azure Storage.",
    "",
    "Parámetros:",
    "- prefix (query, opcional): Prefijo para filtrar archivos",
    "",
    "Respuesta exitosa:",
    "",
    "   {",
    "     \"files\": [",
    "       {",
    "         \"name\": \"ejemplo1.pdf\",",
    "         \"size\": 12345,",
    "         \"content_type\": \"application/pdf\",",
    "         \"url\": \"https://cuenta.blob.core.windows.net/contenedor/ejemplo1.pdf\",",
    "         \"created_on\": \"2025-05-26T12:30:45.123456\"",
    "       },",
    "       {",
    "         \"name\": \"ejemplo2.jpg\",",
    "         \"size\": 67890,",
    "         \"content_type\": \"image/jpeg\",",
    "         \"url\": \"https://cuenta.blob.core.windows.net/contenedor/ejemplo2.jpg\",",
    "         \"created_on\": \"2025-05-26T12:35:12.654321\"",
    "       }",
    "     ],",
    "     \"count\": 2",
    "   }",
    "",
    "Errores posibles:",
    "- 500 Internal Server Error: Error al listar archivos"
]
for item in list_endpoint:
    if item.startswith("   "):
        p = doc.add_paragraph(style='Codigo')
        p.add_run(item.strip())
    else:
        p = doc.add_paragraph(style='TextoNormal')
        p.add_run(item)

# Endpoint: Eliminación de Archivos
doc.add_paragraph('4.2.4 Eliminación de Archivos', style='Subtitulo2')
delete_endpoint = [
    "Endpoint: DELETE /files/{file_name}",
    "",
    "Descripción: Elimina un archivo de Azure Storage.",
    "",
    "Parámetros:",
    "- file_name (path): Nombre del archivo a eliminar",
    "",
    "Respuesta exitosa (204 No Content):",
    "- Sin contenido",
    "",
    "Errores posibles:",
    "- 404 Not Found: El archivo no existe",
    "- 500 Internal Server Error: Error al eliminar el archivo"
]
for item in delete_endpoint:
    p = doc.add_paragraph(style='TextoNormal')
    p.add_run(item)

doc.add_paragraph('4.3 Códigos de Error', style='Subtitulo2')
error_codes = [
    "La API utiliza los siguientes códigos de estado HTTP:",
    "",
    "- 200 OK: La solicitud se ha completado correctamente",
    "- 201 Created: El recurso se ha creado correctamente",
    "- 204 No Content: La solicitud se ha completado correctamente, pero no hay contenido para devolver",
    "- 400 Bad Request: La solicitud es incorrecta o malformada",
    "- 404 Not Found: El recurso solicitado no existe",
    "- 500 Internal Server Error: Error interno del servidor",
    "",
    "Formato de Respuesta de Error:",
    "",
    "   {",
    "     \"error\": \"Mensaje de error\",",
    "     \"detail\": \"Detalles adicionales (opcional)\"",
    "   }"
]
for item in error_codes:
    if item.startswith("   "):
        p = doc.add_paragraph(style='Codigo')
        p.add_run(item.strip())
    else:
        p = doc.add_paragraph(style='TextoNormal')
        p.add_run(item)

# 5. Guía de Usuario
doc.add_paragraph('5. Guía de Usuario', style='Subtitulo1')

doc.add_paragraph('5.1 Uso de la API', style='Subtitulo2')
api_usage = [
    "La API proporciona endpoints para gestionar archivos en Azure Storage:"
]
for item in api_usage:
    p = doc.add_paragraph(style='TextoNormal')
    p.add_run(item)

doc.add_paragraph('5.1.1 Carga de archivos', style='Subtitulo2')
upload_usage = [
    "Para cargar un archivo a Azure Storage:",
    "",
    "1. Vaya a http://localhost:8000/docs en su navegador",
    "2. Expanda el endpoint POST /api/files/upload",
    "3. Haga clic en \"Try it out\"",
    "4. Seleccione un archivo usando el botón \"Choose File\"",
    "5. Haga clic en \"Execute\"",
    "",
    "Alternativamente, puede usar cURL:",
    "",
    "   curl -X POST \"http://localhost:8000/api/files/upload\" \\",
    "     -H \"accept: application/json\" \\",
    "     -H \"Content-Type: multipart/form-data\" \\",
    "     -F \"file=@/ruta/al/archivo.pdf\""
]
for item in upload_usage:
    if item.startswith("   "):
        p = doc.add_paragraph(style='Codigo')
        p.add_run(item.strip())
    else:
        p = doc.add_paragraph(style='TextoNormal')
        p.add_run(item)

doc.add_paragraph('5.1.2 Listado de archivos', style='Subtitulo2')
list_usage = [
    "Para listar los archivos disponibles:",
    "",
    "1. Vaya a http://localhost:8000/docs en su navegador",
    "2. Expanda el endpoint GET /api/files/list",
    "3. Haga clic en \"Try it out\"",
    "4. Opcionalmente, especifique un prefijo para filtrar archivos",
    "5. Haga clic en \"Execute\"",
    "",
    "Alternativamente, puede usar cURL:",
    "",
    "   curl -X GET \"http://localhost:8000/api/files/list\" \\",
    "     -H \"accept: application/json\""
]
for item in list_usage:
    if item.startswith("   "):
        p = doc.add_paragraph(style='Codigo')
        p.add_run(item.strip())
    else:
        p = doc.add_paragraph(style='TextoNormal')
        p.add_run(item)

doc.add_paragraph('5.1.3 Descarga de archivos', style='Subtitulo2')
download_usage = [
    "Para descargar un archivo:",
    "",
    "1. Vaya a http://localhost:8000/docs en su navegador",
    "2. Expanda el endpoint GET /api/files/download/{file_name}",
    "3. Haga clic en \"Try it out\"",
    "4. Introduzca el nombre del archivo a descargar",
    "5. Haga clic en \"Execute\"",
    "",
    "Alternativamente, puede usar cURL:",
    "",
    "   curl -X GET \"http://localhost:8000/api/files/download/nombre_archivo.pdf\" \\",
    "     -H \"accept: application/octet-stream\" \\",
    "     --output archivo_descargado.pdf"
]
for item in download_usage:
    if item.startswith("   "):
        p = doc.add_paragraph(style='Codigo')
        p.add_run(item.strip())
    else:
        p = doc.add_paragraph(style='TextoNormal')
        p.add_run(item)

doc.add_paragraph('5.1.4 Eliminación de archivos', style='Subtitulo2')
delete_usage = [
    "Para eliminar un archivo:",
    "",
    "1. Vaya a http://localhost:8000/docs en su navegador",
    "2. Expanda el endpoint DELETE /api/files/{file_name}",
    "3. Haga clic en \"Try it out\"",
    "4. Introduzca el nombre del archivo a eliminar",
    "5. Haga clic en \"Execute\"",
    "",
    "Alternativamente, puede usar cURL:",
    "",
    "   curl -X DELETE \"http://localhost:8000/api/files/nombre_archivo.pdf\" \\",
    "     -H \"accept: application/json\""
]
for item in delete_usage:
    if item.startswith("   "):
        p = doc.add_paragraph(style='Codigo')
        p.add_run(item.strip())
    else:
        p = doc.add_paragraph(style='TextoNormal')
        p.add_run(item)

doc.add_paragraph('5.2 Ejemplos prácticos', style='Subtitulo2')

doc.add_paragraph('5.2.1 Ejemplo: Carga y descarga de un archivo', style='Subtitulo2')
example1 = [
    "   # Cargar un archivo",
    "   curl -X POST \"http://localhost:8000/api/files/upload\" \\",
    "     -H \"accept: application/json\" \\",
    "     -H \"Content-Type: multipart/form-data\" \\",
    "     -F \"file=@/ruta/al/documento.pdf\"",
    "",
    "   # La respuesta incluirá el nombre del archivo en Azure Storage, por ejemplo:",
    "   # {\"file_name\":\"a1b2c3d4_documento.pdf\",\"file_size\":12345,\"content_type\":\"application/pdf\",\"url\":\"https://...\"}",
    "",
    "   # Descargar el archivo usando el nombre devuelto",
    "   curl -X GET \"http://localhost:8000/api/files/download/a1b2c3d4_documento.pdf\" \\",
    "     -H \"accept: application/octet-stream\" \\",
    "     --output documento_descargado.pdf"
]
for item in example1:
    p = doc.add_paragraph(style='Codigo')
    p.add_run(item.strip())

doc.add_paragraph('5.2.2 Ejemplo: Listar y eliminar archivos', style='Subtitulo2')
example2 = [
    "   # Listar todos los archivos",
    "   curl -X GET \"http://localhost:8000/api/files/list\" \\",
    "     -H \"accept: application/json\"",
    "",
    "   # Listar archivos con un prefijo específico",
    "   curl -X GET \"http://localhost:8000/api/files/list?prefix=documento\" \\",
    "     -H \"accept: application/json\"",
    "",
    "   # Eliminar un archivo específico",
    "   curl -X DELETE \"http://localhost:8000/api/files/a1b2c3d4_documento.pdf\" \\",
    "     -H \"accept: application/json\""
]
for item in example2:
    p = doc.add_paragraph(style='Codigo')
    p.add_run(item.strip())

doc.add_paragraph('5.3 Solución de problemas', style='Subtitulo2')
troubleshooting = [
    "Problemas comunes:",
    "",
    "1. Error de conexión a Azure Storage:",
    "   - Verifique que las credenciales en el archivo .env sean correctas",
    "   - Asegúrese de que su cuenta de Azure Storage esté activa",
    "   - Compruebe que tiene permisos suficientes para las operaciones",
    "",
    "2. Error al iniciar la aplicación:",
    "   - Verifique que todas las dependencias estén instaladas: pip install -r requirements.txt",
    "   - Asegúrese de que el puerto 8000 no esté en uso por otra aplicación",
    "",
    "3. Error al cargar archivos grandes:",
    "   - La API tiene un límite de tamaño para la carga de archivos",
    "   - Para archivos muy grandes, considere dividirlos o ajustar la configuración de FastAPI",
    "",
    "Para obtener más información sobre los errores, puede iniciar la aplicación en modo debug:",
    "",
    "   # Edite el archivo .env y establezca:",
    "   # DEBUG=True",
    "",
    "   # Luego inicie la aplicación:",
    "   uvicorn src.main:app --reload --host 0.0.0.0 --port 8000"
]
for item in troubleshooting:
    if item.startswith("   "):
        p = doc.add_paragraph(style='Codigo')
        p.add_run(item.strip())
    else:
        p = doc.add_paragraph(style='TextoNormal')
        p.add_run(item)

# 6. Estructura del Proyecto
doc.add_paragraph('6. Estructura del Proyecto', style='Subtitulo1')
project_structure = [
    "A continuación se muestra la estructura completa del proyecto:",
    "",
    "   azure_storage_app/",
    "   ├── src/                          # Código fuente de la aplicación",
    "   │   ├── __init__.py               # Inicializador del paquete",
    "   │   ├── main.py                   # Punto de entrada de la aplicación",
    "   │   ├── config/                   # Configuración de la aplicación",
    "   │   │   ├── __init__.py",
    "   │   │   ├── settings.py           # Configuración general y variables de entorno",
    "   │   │   └── azure_config.py       # Configuración específica de Azure Storage",
    "   │   │",
    "   │   ├── models/                   # Modelos de datos y esquemas",
    "   │   │   ├── __init__.py",
    "   │   │   └── schemas.py            # Esquemas Pydantic para validación de datos",
    "   │   │",
    "   │   ├── services/                 # Servicios de negocio",
    "   │   │   ├── __init__.py",
    "   │   │   └── storage_service.py    # Servicio para operaciones con Azure Storage",
    "   │   │",
    "   │   ├── api/                      # Endpoints de la API",
    "   │   │   ├── __init__.py",
    "   │   │   ├── routes.py             # Registro de rutas",
    "   │   │   └── endpoints/            # Endpoints específicos",
    "   │   │       ├── __init__.py",
    "   │   │       └── files.py          # Endpoints para gestión de archivos",
    "   │   │",
    "   │   └── utils/                    # Utilidades y helpers",
    "   │       ├── __init__.py",
    "   │       ├── exceptions.py         # Excepciones personalizadas",
    "   │       └── helpers.py            # Funciones auxiliares",
    "   │",
    "   ├── tests/                        # Pruebas unitarias y de integración",
    "   │   ├── __init__.py",
    "   │   ├── test_api.py               # Pruebas para la API",
    "   │   └── test_services.py          # Pruebas para los servicios",
    "   │",
    "   ├── docs/                         # Documentación",
    "   │   ├── requerimientos.md         # Análisis de requerimientos",
    "   │   ├── arquitectura.md           # Diseño de arquitectura",
    "   │   ├── api.md                    # Documentación de la API",
    "   │   └── guia_usuario.md           # Guía de usuario",
    "   │",
    "   ├── .env.example                  # Ejemplo de variables de entorno",
    "   ├── requirements.txt              # Dependencias del proyecto",
    "   └── README.md                     # Documentación principal"
]
for item in project_structure:
    if item.startswith("   "):
        p = doc.add_paragraph(style='Codigo')
        p.add_run(item.strip())
    else:
        p = doc.add_paragraph(style='TextoNormal')
        p.add_run(item)

# Guardar el documento
doc.save('/home/ubuntu/azure_storage_app/docs/Documentacion_Azure_Storage_App.docx')

print("Documento Word creado exitosamente en: /home/ubuntu/azure_storage_app/docs/Documentacion_Azure_Storage_App.docx")
